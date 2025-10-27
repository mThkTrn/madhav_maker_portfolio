# modules/docx_to_txt.py
# Converts a docx file to a txt file while annotating bold, italic, and underlined parts.

import docx
import sys

def parse_paragraph(paragraph: docx.text.paragraph.Paragraph) -> str:
    text = ""
    for runs in paragraph.runs:
        if len(runs.text.strip()) == 0:
            text += runs.text
            continue

        run = runs.text

        if runs.underline:
            run = "{u}" + run + "{/u}"
        if runs.bold:
            run = "{b}" + run + "{/b}"
        if runs.italic:
            run = "{i}" + run + "{/i}"

        text += run
    return text.strip()

def main(input_file, output_file):
    doc = docx.Document(input_file)
    with open(output_file, "w", encoding="utf-8") as f:
        # Process all paragraphs in the document
        for para in doc.paragraphs:
            line = parse_paragraph(para)
            if line:
                f.write(line + "\n")
                
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        line = parse_paragraph(para)
                        if line:
                            f.write(line + "\n")

def cli():
    """Command line interface for the script"""
    if len(sys.argv) != 3:
        print("Usage: python docx_to_txt.py input.docx output.txt")
        sys.exit(1)
    main(sys.argv[1], sys.argv[2])

if __name__ == "__main__":
    cli()
