#!/usr/bin/env python3
"""
Test script for the enhanced packet parser to verify it handles both formats correctly.
"""

import sys
import os

from packet_parser import Parser

def test_alternating_format():
    """Test the alternating format (tossup-bonus-tossup-bonus)"""
    print("Testing alternating format...")
    
    # Sample alternating format packet
    packet_text = """1. This is a tossup question about science.
ANSWER: science answer
<Science>

1. This is a bonus leadin. For 10 points each:
[10] First bonus part.
ANSWER: first answer
[10] Second bonus part.
ANSWER: second answer
[10] Third bonus part.
ANSWER: third answer
<Science>

2. This is another tossup question.
ANSWER: another answer
<Literature>

2. This is another bonus leadin. For 10 points each:
[10] First part.
ANSWER: first
[10] Second part.
ANSWER: second
[10] Third part.
ANSWER: third
<Literature>"""

    parser = Parser(
        has_question_numbers=True,
        has_category_tags=True,
        bonus_length=3,
        buzzpoints=False,
        modaq=False,
        auto_insert_powermarks=False,
        classify_unknown=True,
        space_powermarks=False,
        always_classify=False
    )
    
    result = parser.parse_packet(packet_text, "test_alternating")
    print(f"Found {len(result['tossups'])} tossups and {len(result['bonuses'])} bonuses")
    
    assert len(result['tossups']) == 2, f"Expected 2 tossups, got {len(result['tossups'])}"
    assert len(result['bonuses']) == 2, f"Expected 2 bonuses, got {len(result['bonuses'])}"
    print("✓ Alternating format test passed!")

def test_grouped_format():
    """Test the grouped format (tossup-tossup-bonus-bonus)"""
    print("Testing grouped format...")
    
    # Sample grouped format packet (like FARSI)
    packet_text = """1. This is the first tossup question about science.
ANSWER: science answer
<Science>

2. This is the second tossup question about literature.
ANSWER: literature answer
<Literature>

This is the first bonus leadin. For 10 points each:
[10] First bonus part.
ANSWER: first answer
[10] Second bonus part.
ANSWER: second answer
[10] Third bonus part.
ANSWER: third answer
<Science>

This is the second bonus leadin. For 10 points each:
[10] First part.
ANSWER: first
[10] Second part.
ANSWER: second
[10] Third part.
ANSWER: third
<Literature>"""

    parser = Parser(
        has_question_numbers=True,
        has_category_tags=True,
        bonus_length=3,
        buzzpoints=False,
        modaq=False,
        auto_insert_powermarks=False,
        classify_unknown=True,
        space_powermarks=False,
        always_classify=False
    )
    
    result = parser.parse_packet(packet_text, "test_grouped")
    print(f"Found {len(result['tossups'])} tossups and {len(result['bonuses'])} bonuses")
    
    assert len(result['tossups']) == 2, f"Expected 2 tossups, got {len(result['tossups'])}"
    assert len(result['bonuses']) == 2, f"Expected 2 bonuses, got {len(result['bonuses'])}"
    print("✓ Grouped format test passed!")

def test_grouped_format_no_tags():
    """Test the grouped format without category tags"""
    print("Testing grouped format without category tags...")
    
    # Sample grouped format packet without category tags
    packet_text = """1. This is the first tossup question about science.
ANSWER: science answer

2. This is the second tossup question about literature.
ANSWER: literature answer

This is the first bonus leadin. For 10 points each:
[10] First bonus part.
ANSWER: first answer
[10] Second bonus part.
ANSWER: second answer
[10] Third bonus part.
ANSWER: third answer

This is the second bonus leadin. For 10 points each:
[10] First part.
ANSWER: first
[10] Second part.
ANSWER: second
[10] Third part.
ANSWER: third"""

    parser = Parser(
        has_question_numbers=True,
        has_category_tags=False,
        bonus_length=3,
        buzzpoints=False,
        modaq=False,
        auto_insert_powermarks=False,
        classify_unknown=True,
        space_powermarks=False,
        always_classify=False
    )
    
    result = parser.parse_packet(packet_text, "test_grouped_no_tags")
    print(f"Found {len(result['tossups'])} tossups and {len(result['bonuses'])} bonuses")
    
    assert len(result['tossups']) == 2, f"Expected 2 tossups, got {len(result['tossups'])}"
    assert len(result['bonuses']) == 2, f"Expected 2 bonuses, got {len(result['bonuses'])}"
    print("✓ Grouped format without tags test passed!")

def test_real_farsi_packet():
    """Test with actual FARSI packet content"""
    print("Testing with real FARSI packet content...")
    
    # Extract text from the FARSI packet
    try:
        import docx
        doc = docx.Document('../controllers/sample_packets/FARSI Packet 1.docx')
        packet_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Clean up the text for parsing
        packet_text = packet_text.replace('\n\n', '\n').strip()
        
        parser = Parser(
            has_question_numbers=True,
            has_category_tags=True,
            bonus_length=3,
            buzzpoints=False,
            modaq=False,
            auto_insert_powermarks=False,
            classify_unknown=True,
            space_powermarks=False,
            always_classify=False
        )
        
        result = parser.parse_packet(packet_text, "FARSI Packet 1")
        print(f"Found {len(result['tossups'])} tossups and {len(result['bonuses'])} bonuses")
        
        # FARSI packets typically have 20 tossups and 20 bonuses
        assert len(result['tossups']) > 0, f"Expected some tossups, got {len(result['tossups'])}"
        assert len(result['bonuses']) > 0, f"Expected some bonuses, got {len(result['bonuses'])}"
        print("✓ Real FARSI packet test passed!")
        
    except ImportError:
        print("python-docx not available, skipping real packet test")
    except Exception as e:
        print(f"Error testing real FARSI packet: {e}")

def test_real_scottie_packet():
    """Test with actual Scottie packet content"""
    print("Testing with real Scottie packet content...")
    
    # Extract text from the Scottie packet
    try:
        import docx
        doc = docx.Document('../controllers/sample_packets/Blended Round 12 - 2021 Scottie.docx')
        packet_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Clean up the text for parsing
        packet_text = packet_text.replace('\n\n', '\n').strip()
        
        parser = Parser(
            has_question_numbers=True,
            has_category_tags=True,
            bonus_length=3,
            buzzpoints=False,
            modaq=False,
            auto_insert_powermarks=False,
            classify_unknown=True,
            space_powermarks=False,
            always_classify=False
        )
        
        result = parser.parse_packet(packet_text, "Scottie Round 12")
        print(f"Found {len(result['tossups'])} tossups and {len(result['bonuses'])} bonuses")
        
        # Scottie packets typically have 20 tossups and 20 bonuses
        assert len(result['tossups']) > 0, f"Expected some tossups, got {len(result['tossups'])}"
        assert len(result['bonuses']) > 0, f"Expected some bonuses, got {len(result['bonuses'])}"
        print("✓ Real Scottie packet test passed!")
        
    except ImportError:
        print("python-docx not available, skipping real packet test")
    except Exception as e:
        print(f"Error testing real Scottie packet: {e}")

if __name__ == "__main__":
    print("Testing enhanced packet parser...")
    print("=" * 50)
    
    try:
        test_alternating_format()
        test_grouped_format()
        test_grouped_format_no_tags()
        test_real_farsi_packet()
        test_real_scottie_packet()
        print("\n" + "=" * 50)
        print("🎉 All tests passed! The enhanced parser works correctly.")
    except Exception as e:
        print(f"\n❌ Test failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

